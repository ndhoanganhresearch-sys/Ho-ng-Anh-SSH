"""Convert PAPER_DRAFT.md to a formatted Word document (IEEE-style)."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, os

PAPER_PATH = os.path.join(os.path.dirname(__file__), "PAPER_DRAFT.md")
OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "PAPER_SSL_Tunnel.docx")

def set_font(run, name="Times New Roman", size=10, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12 if level == 1 else 11)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_paragraph(doc, text, indent=False, code=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    if code:
        p.paragraph_format.left_indent = Inches(0.4)
        run = p.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        return p
    # Handle **bold** and *italic*
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_font(run, bold=True)
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            set_font(run, italic=True)
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            run = p.add_run(part)
            set_font(run)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    return p

def add_table(doc, rows_data, headers=None):
    cols = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data) + (1 if headers else 0), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    start_row = 0
    if headers:
        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = h
            run = cell.paragraphs[0].runs[0]
            set_font(run, bold=True, size=9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        start_row = 1

    for i, row in enumerate(rows_data):
        for j, val in enumerate(row):
            cell = table.rows[i + start_row].cells[j]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            set_font(run, size=9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    return table

def build_docx():
    doc = Document()

    # Page margins (narrow for IEEE style)
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Title ──
    title = doc.add_heading("SSL Smart Tunnel Monitoring System: An Automated LiDAR-Based Point Cloud Processing Pipeline for Structural Health Assessment", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

    # Author
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Nguyen Duy Hoang Anh")
    set_font(run, bold=True, size=11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[University / Department]  |  June 2026")
    set_font(run, italic=True, size=10)
    doc.add_paragraph()

    # Read markdown
    with open(PAPER_PATH, encoding="utf-8") as f:
        lines = f.readlines()

    in_code = False
    code_buf = []
    in_table = False
    table_buf = []

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Code block
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                for cl in code_buf:
                    add_paragraph(doc, cl, code=True)
                doc.add_paragraph()
                code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Table
        if "|" in line and line.strip().startswith("|"):
            if not in_table:
                in_table = True
                table_buf = []
            # Skip separator rows
            if re.match(r"^\|[-| :]+\|$", line.strip()):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_buf.append(cells)
            i += 1
            continue
        else:
            if in_table and table_buf:
                in_table = False
                add_table(doc, table_buf[1:], headers=table_buf[0])
                table_buf = []

        # Skip the title line (already added)
        if line.startswith("# SSL Smart"):
            i += 1
            continue

        # Section headings
        if line.startswith("## "):
            add_heading(doc, line[3:], level=1)
            i += 1
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:], level=2)
            i += 1
            continue

        # Author/date lines (skip, already added)
        if line.startswith("**Author:**") or line.startswith("**Affiliation:**") or line.startswith("**Date:**"):
            i += 1
            continue

        # Horizontal rule
        if line.startswith("---"):
            doc.add_paragraph().add_run("─" * 60).font.size = Pt(6)
            i += 1
            continue

        # Bullet list
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.3)
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', line[2:])
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_font(run, bold=True)
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                else:
                    run = p.add_run(part)
                    set_font(run)
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(re.sub(r"^\d+\. ", "", line))
            set_font(run)
            i += 1
            continue

        # Normal paragraph
        if line.strip():
            add_paragraph(doc, line.strip())

        i += 1

    doc.save(OUTPUT_PATH)
    print(f"[OK] Saved: {OUTPUT_PATH}")

if __name__ == "__main__":
    build_docx()
