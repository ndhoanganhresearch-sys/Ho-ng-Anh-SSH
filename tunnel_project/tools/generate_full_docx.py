"""Render the assembled SSL Tunnel manuscript Markdown to a formatted DOCX.

Input : paper/drafts/main_paper_full_assembled.md
Output: paper/drafts/SSL_Tunnel_Full_Paper_v1.docx

Formatting follows paper/reviews/WRITING_RULES.md Section 9:
Times New Roman, title 16pt, headings 14pt, body 12pt, 1.5 line spacing,
US Letter, 1-inch margins.

Handles: title (#), headings (##, ###), paragraphs with **bold**/*italic*,
markdown pipe tables, display equations ($$...$$), and bold table/figure captions.
LaTeX math is converted to a readable Unicode approximation (re-typeset in Word's
equation editor for camera-ready).

Run from tunnel_project/:
    ..\\.venv\\Scripts\\python.exe tools\\generate_full_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

HERE = Path(__file__).resolve().parent
DRAFTS = HERE.parent / "paper" / "drafts"
SRC = DRAFTS / "main_paper_full_assembled.md"
OUT = DRAFTS / "SSL_Tunnel_Full_Paper_v1.docx"

FONT = "Times New Roman"
SZ_TITLE, SZ_H1, SZ_H2, SZ_BODY = 16, 14, 13, 12

# --- LaTeX -> Unicode approximation -----------------------------------------
_GREEK = {
    r"\\sigma": "σ", r"\\lambda": "λ", r"\\varepsilon": "ε", r"\\tau": "τ",
    r"\\delta": "δ", r"\\Delta": "Δ", r"\\hat\{z\}": "ẑ", r"\\hat\{c\}": "ĉ",
    r"\\sqrt": "√", r"\\cdot": "·", r"\\times": "×", r"\\le": "≤", r"\\ge": "≥",
    r"\\iff": "⟺", r"\\max": "max", r"\\min": "min", r"\\mathrm": "",
    r"\\qquad": "    ", r"\\quad": "  ", r"\\lVert": "‖", r"\\rVert": "‖",
    r"\\big": "", r"\\!": "", r"\\,": " ", r"\\;": " ", r"\\text": "",
}


def latex_to_text(s: str) -> str:
    s = s.strip()
    s = re.sub(r"\\tag\{([^}]*)\}", r"  (\1)", s)
    s = re.sub(r"\\frac\{([^{}]*)\}\{([^{}]*)\}", r"(\1)/(\2)", s)
    for pat, rep in _GREEK.items():
        s = re.sub(pat, rep, s)
    s = re.sub(r"_\{([^}]*)\}", r"_\1", s)   # keep subscripts inline
    s = re.sub(r"\^\{([^}]*)\}", r"^\1", s)
    s = s.replace("{", "").replace("}", "")
    s = s.replace("\\", "")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def clean_inline(s: str) -> str:
    """Convert inline $..$ math and strip simple markdown emphasis markers."""
    s = re.sub(r"\$([^$]*)\$", lambda m: latex_to_text(m.group(1)), s)
    return s


_BOLD = re.compile(r"\*\*(.+?)\*\*")
_ITAL = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")


def add_runs(paragraph, text: str) -> None:
    """Add text to a paragraph honouring **bold** and *italic* segments."""
    text = clean_inline(text)
    # Tokenise into (style, chunk) by scanning bold first, then italic.
    pos = 0
    tokens: list[tuple[str, str]] = []
    for m in _BOLD.finditer(text):
        if m.start() > pos:
            tokens.append(("plain", text[pos:m.start()]))
        tokens.append(("bold", m.group(1)))
        pos = m.end()
    if pos < len(text):
        tokens.append(("plain", text[pos:]))

    for style, chunk in tokens:
        if style == "bold":
            r = paragraph.add_run(chunk)
            r.bold = True
            _font(r)
            continue
        # italic within plain chunk
        ip = 0
        for im in _ITAL.finditer(chunk):
            if im.start() > ip:
                _font(paragraph.add_run(chunk[ip:im.start()]))
            r = paragraph.add_run(im.group(1))
            r.italic = True
            _font(r)
            ip = im.end()
        if ip < len(chunk):
            _font(paragraph.add_run(chunk[ip:]))


def _font(run, size: int = SZ_BODY) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)


def _spacing(p) -> None:
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(8)


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(SZ_BODY)

    for section in doc.sections:
        section.page_width = Pt(612)      # US Letter 8.5in
        section.page_height = Pt(792)     # 11in
        for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
            setattr(section, m, Pt(72))   # 1 inch

    lines = SRC.read_text(encoding="utf-8").splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Title
        if stripped.startswith("# ") and not stripped.startswith("## "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(stripped[2:].strip())
            r.bold = True
            _font(r, SZ_TITLE)
            _spacing(p)
            i += 1
            continue

        # Headings
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[4:].strip())
            r.bold = True
            _font(r, SZ_H2)
            _spacing(p)
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[3:].strip())
            r.bold = True
            _font(r, SZ_H1)
            _spacing(p)
            i += 1
            continue

        # Display equation
        if stripped.startswith("$$"):
            eq = stripped.strip("$ ").strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(latex_to_text(eq))
            r.italic = True
            _font(r)
            _spacing(p)
            i += 1
            continue

        # Markdown table block
        if stripped.startswith("|") and i + 1 < n and set(lines[i + 1].strip()) <= set("|-: "):
            header = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2  # skip header + separator
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, htext in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.paragraphs[0].text = ""
                add_runs(cell.paragraphs[0], htext)
                for rr in cell.paragraphs[0].runs:
                    rr.bold = True
            for row in rows:
                cells = table.add_row().cells
                for j, ctext in enumerate(row[:len(header)]):
                    cells[j].paragraphs[0].text = ""
                    add_runs(cells[j].paragraphs[0], ctext)
                    for rr in cells[j].paragraphs[0].runs:
                        rr.font.size = Pt(11)
            doc.add_paragraph()
            continue

        # Numbered contribution / ordinary paragraph
        p = doc.add_paragraph()
        add_runs(p, stripped)
        _spacing(p)
        i += 1

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
