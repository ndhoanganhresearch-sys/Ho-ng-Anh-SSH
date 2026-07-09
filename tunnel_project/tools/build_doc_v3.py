#!/usr/bin/env python3
"""
Rewrite the 2-3 page technical document for Prof. Yoon so that Section 3
(Demo Video) reflects the ACTUAL recorded demo of the tool (one-click Auto
Pipeline over T0~T5), instead of the previous hypothetical scenario.

Sources of truth for numbers:
  output/timeseries_benchmark/timeseries_benchmark_report.json
  CLAUDE.md ground-truth spec
  docs/Ghi man hinh 3 ... .mp4  (real screen recording -> demo walkthrough)
  docs/LiDAR_Tunnel_RealDemo_Annotated.mp4 (annotated deliverable)
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DOCS = Path(r"C:\Users\ssl\Desktop\Code Python\data python cusor\docs")
OUT = DOCS / "LiDAR_Tunnel_TimeSeries_Analysis_v4.docx"

NAVY = RGBColor(0x1F, 0x3C, 0x72)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# base style
normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10.5)

for i in range(1, 4):
    st = doc.styles[f"Heading {i}"]
    st.font.name = "Arial"
    st.font.color.rgb = NAVY
    st.font.size = Pt({1: 14, 2: 12, 3: 11}[i])
    st.font.bold = True


def title(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "Arial"
    r.font.color.rgb = NAVY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def sub(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = GREY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def body(text):
    return doc.add_paragraph(text)


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def kv_bullet(key, val):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(key + ": ")
    r.bold = True
    p.add_run(val)
    return p


# ---------------------------------------------------------------- header
title("LiDAR-Based Tunnel Time-Series Shape Analysis Technology")
sub("SSL Smart Tunnel Monitoring System  |  Step 6 — Time-Series Analysis Module")

# ---------------------------------------------------------------- 1
doc.add_heading("1. Technical Overview and Necessity", level=1)
body(
    "Railway tunnels are long-lived linear structures in which shape changes — crown "
    "settlement, sidewall convergence, sectional eccentricity, and local damage — "
    "accumulate over years of operation. Conventional visual, tape, and total-station "
    "inspection is operator-dependent and typically captures only a single qualitative "
    "snapshot. What is needed is a time-series framework that quantitatively tracks how "
    "the tunnel shape evolves and objectively evaluates the location, magnitude, and rate "
    "of deformation."
)
body(
    "Terrestrial LiDAR acquires dense 3D point clouds of the tunnel interior contact-free "
    "and at millimetre precision, even under live-traffic or maintenance conditions, making "
    "it an ideal input for digital-twin-based maintenance. The SSL Smart Tunnel Monitoring "
    "System is a fully automated Python pipeline that co-registers LiDAR scans "
    "from different epochs into a common reference frame, extracts shape change relative to "
    "the tunnel centreline and cross-sections, and combines M3C2-based 4D change detection "
    "with per-section engineering indices, producing report-ready outputs for "
    "digital-twin-based maintenance."
)
body(
    "This technology corresponds to the mid-term consulting item “Intelligent Analysis–"
    "Simulation Technology 1: LiDAR-Based Tunnel Time-Series Shape Analysis” "
    "(DT 중간컨설팅_260707.pptx, slide 39) and is implemented as Step 6 of the tool, "
    "taking a T0 reference scan and Tn comparison scans — or a full T0~T5 multi-epoch set — "
    "as input."
)

# ---------------------------------------------------------------- 2
doc.add_heading("2. Core Research Content", level=1)

doc.add_heading("2.1 Multi-Point LiDAR Registration", level=2)
body(
    "Time-series comparison requires every epoch to share one reference coordinate frame; "
    "sensor pose, trajectory, and IMU error otherwise mask real structural change. T0 is "
    "fixed as the reference epoch and all other epochs are aligned to it."
)
kv_bullet("Coarse alignment", "feature/target-based initial transform (GROR + FPFH)")
kv_bullet("Fine registration", "Generalized ICP (small_gicp), Open3D point-to-plane ICP as fallback")
kv_bullet("Benchmark", "0.196 mm RMSE, 339 ms (separate registration benchmark) — faster and more stable than Open3D ICP")
kv_bullet("Quality gate", "registration RMSE and residuals checked; high RMSE flags results as low-confidence")

doc.add_heading("2.2 Geometric Modeling and Shape Indices", level=2)
body(
    "The tunnel centreline is estimated and Frenet-frame cross-sections orthogonal to it are "
    "generated per chainage, which reduces apparent-ovality error on curved tunnels. Each "
    "section is compared with the design radius or the T0 reference section."
)
kv_bullet("Shape indices", "crown settlement, lateral convergence, ovality, eccentricity, local damage")
kv_bullet("Section fitting", "Fitzgibbon Direct Least-Squares ellipse fitting for precise section geometry")
kv_bullet("Output", "per-section tables and graphs; over-threshold zones highlighted in the 3D heatmap / 2D overlay")

doc.add_heading("2.3 M3C2-Based 4D Change Detection and Trend Analysis", level=2)
body(
    "Using py4dgeo, T0 core points and local normals are fixed and the signed distance to "
    "each epoch is computed along the surface normal. A Level-of-Detection (LoD) criterion "
    "separates statistically significant deformation from noise, and results are aggregated "
    "by chainage, section angle, and epoch to form trends and a forecast."
)
kv_bullet("Validation dataset", "80 m length, 3.0 m radius, 6 epochs (T0~T5), 15,456 points/epoch")
kv_bullet("Ground truth @ T5", "crown -45 mm @20 m, convergence -35 mm @45 m, local damage -40 mm @65 m (from T3)")
kv_bullet("Measured @ T5", "crown max 44.05 mm, heatmap max 45.0 mm, ovality 0.20 %, eccentricity 1.52 mm")
kv_bullet("Forecast", "linear-trend R² ≈ 0.9997; caution at ~9 epochs, critical at ~15 epochs")

# ---------------------------------------------------------------- 3 (rewritten)
doc.add_heading("3. Demo Video Walkthrough", level=1)
body(
    "The demo is a real screen recording of the SSL tool running the one-click Auto Pipeline "
    "on the validated T0~T5 dataset (file: docs/LiDAR_Tunnel_RealDemo_Annotated.mp4, "
    "~1 minute, with on-screen step captions). Its purpose is to show how raw LiDAR point "
    "clouds become maintenance decision information, not to explain the algorithms. After the "
    "six epochs are loaded, a single “Auto Pipeline (1-click full analysis)” action runs "
    "the entire Step 1–Step 6 chain, and the following panels appear in sequence:"
)

tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Light Grid Accent 1"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for c, t in zip(hdr, ["Time", "On-screen panel", "What it demonstrates"]):
    run = c.paragraphs[0].add_run(t)
    run.bold = True
    run.font.size = Pt(9.5)

rows = [
    ("0:03", "Auto Pipeline + Structure tree (S1–S6)",
     "Six epochs T0~T5 loaded; one click runs preprocess → register → sections → M3C2."),
    ("0:15", "2D Cross-Section (circle fit)",
     "Per-chainage section with crown/wall/floor deviation, ovality and eccentricity."),
    ("0:28", "Multi-Epoch Deformation Trend",
     "p95/median per epoch with Safe/Caution/Critical bands and a per-epoch status table."),
    ("0:36", "M3C2 Deformation Map (T0→Tn)",
     "Signed surface change with LoD; mm colorbar; results log of all extracted parameters."),
    ("0:45", "Plot 2D Technical Section",
     "Engineering section vs vehicle clearance; exportable to PDF / IFC."),
    ("0:56", "Summary Dashboard",
     "Critical zones, max deformation, trend forecast, and digital-twin handoff."),
]
for time, panel, what in rows:
    cells = tbl.add_row().cells
    for cell, txt, sz in zip(cells, (time, panel, what), (9, 9, 9)):
        r = cell.paragraphs[0].add_run(txt)
        r.font.size = Pt(sz)

body("")  # spacer

# ---------------------------------------------------------------- 4
doc.add_heading("4. Research Outcomes and Future Improvements", level=1)
p = doc.add_paragraph()
p.add_run("Outcomes").bold = True
bullet("End-to-end automated pipeline reducing a manual multi-step workflow to a one-click run.")
bullet("Centreline + Frenet sectioning with five engineering shape indices per section.")
bullet("Benchmark-validated M3C2 tracking on T0~T5 (crown max 44.05 mm vs -45 mm ground truth; forecast R² ≈ 0.9997).")
bullet("Full IFC4X3 export for BIM interoperability with digital-twin platforms.")
bullet("Integrated RAG-LLM / AI Engineering Assistant for decision support.")

p = doc.add_paragraph()
p.add_run("Future Improvements").bold = True
bullet("Robust validation on real field LiDAR (sensor drift, occlusion, reflectivity, cable/equipment clutter).")
bullet("Multi-epoch UI workflow selecting T0~T5 at once with epoch-labelled graphs.")
bullet("Separate cumulative (T0→Tn) and incremental (Tn→Tn+1) deformation in the report.")
bullet("Link LoD significance to structural safety thresholds for clear caution/critical limits.")
bullet("Non-circular profiles (horseshoe), fracture detection, and water-seepage mapping.")

body(
    "With these enhancements, this technology can become a core intelligent analysis–"
    "simulation module for railway-tunnel digital twins."
)

try:
    doc.save(str(OUT))
except PermissionError:
    OUT = OUT.with_name("LiDAR_Tunnel_TimeSeries_Analysis_v5.docx")
    doc.save(str(OUT))
print("[OK] wrote", OUT.name)

# quick page/word estimate
words = sum(len(p.text.split()) for p in doc.paragraphs)
print("paragraphs:", len(doc.paragraphs), "approx words:", words)
