"""Generate PAPER_SSL_Tunnel_V2.docx from PAPER_DRAFT_V2.md in journal style."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re, os

PAPER_PATH = os.path.join(os.path.dirname(__file__), "drafts", "main_paper_working.md")
OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "generated", "SSL_Tunnel_main_paper_working.docx")

def sf(run, name="Times New Roman", size=10, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def render_inline(p, text, italic_all=False):
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); sf(r, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1]); sf(r, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name = "Courier New"; r.font.size = Pt(9)
        else:
            r = p.add_run(part); sf(r, italic=italic_all)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sizes = {1: 12, 2: 11, 3: 10}
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(sizes.get(level, 10))
        run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_para(doc, text, center=False, italic_all=False, code=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if code:
        p.paragraph_format.left_indent = Inches(0.5)
        r = p.add_run(text)
        r.font.name = "Courier New"; r.font.size = Pt(8)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        return p
    render_inline(p, text, italic_all=italic_all)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    return p

def add_table(doc, rows, headers=None):
    if not rows:
        return
    cols = max(len(r) for r in (rows + ([headers] if headers else [])))
    n = len(rows) + (1 if headers else 0)
    t = doc.add_table(rows=n, cols=cols)
    t.style = "Table Grid"
    if headers:
        for j, h in enumerate(headers):
            if j < cols:
                cell = t.rows[0].cells[j]
                cell.text = h
                r = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(h)
                sf(r, bold=True, size=9)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        ri = i + (1 if headers else 0)
        for j, v in enumerate(row):
            if j < cols:
                cell = t.rows[ri].cells[j]
                cell.text = str(v)
                r = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(str(v))
                sf(r, size=9)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def build():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.0)

    with open(PAPER_PATH, encoding="utf-8") as f:
        lines = f.readlines()

    in_code   = False
    code_buf  = []
    in_table  = False
    table_buf = []
    title_done = False

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # ── Code block ──
        if line.startswith("```"):
            if not in_code:
                in_code = True; code_buf = []
            else:
                in_code = False
                for cl in code_buf:
                    add_para(doc, cl, code=True)
                doc.add_paragraph()
                code_buf = []
            i += 1; continue
        if in_code:
            code_buf.append(line); i += 1; continue

        # ── Table ──
        if "|" in line and line.strip().startswith("|"):
            if not in_table:
                in_table = True; table_buf = []
            if re.match(r"^\|[-| :]+\|$", line.strip()):
                i += 1; continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_buf.append(cells); i += 1; continue
        else:
            if in_table and table_buf:
                in_table = False
                add_table(doc, table_buf[1:], headers=table_buf[0])
                table_buf = []

        # ── Title h1 ──
        if line.startswith("# ") and not title_done:
            p = doc.add_heading(line[2:], 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(14)
                r.font.color.rgb = RGBColor(0, 0, 0)
            p.paragraph_format.space_after = Pt(10)
            title_done = True; i += 1; continue

        # ── Author / affiliation / article info ──
        if line.startswith("**Nguyen") or line.startswith("[Department]") or line.startswith("E-mail:"):
            stripped = line.replace("**", "")
            add_para(doc, stripped, center=True, italic_all=("E-mail" in line))
            i += 1; continue

        if line.startswith("**Article Info**") or line.startswith("Received:") or line.startswith("Keywords:"):
            stripped = line.replace("**", "")
            add_para(doc, stripped, italic_all=True)
            i += 1; continue

        # ── Headings ──
        if line.startswith("### "): add_heading(doc, line[4:], 3); i += 1; continue
        if line.startswith("## "):  add_heading(doc, line[3:], 2); i += 1; continue
        if line.startswith("# "):   add_heading(doc, line[2:], 1); i += 1; continue

        # ── Horizontal rule ──
        if line.strip() == "---": doc.add_paragraph(); i += 1; continue

        # ── Bullets ──
        m_bullet = re.match(r"^[-*] (.+)", line)
        if m_bullet:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.3)
            render_inline(p, m_bullet.group(1))
            i += 1; continue

        # ── Numbered list ──
        m_num = re.match(r"^\d+\. (.+)", line)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.3)
            render_inline(p, m_num.group(1))
            i += 1; continue

        # ── Normal paragraph ──
        if line.strip():
            add_para(doc, line.strip())
        else:
            if i > 0 and lines[i-1].strip():
                doc.add_paragraph()

        i += 1

    doc.save(OUTPUT_PATH)
    print("[OK] Saved:", OUTPUT_PATH)

if __name__ == "__main__":
    build()
